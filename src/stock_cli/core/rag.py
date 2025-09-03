"""RAG支持模块，提供基本的嵌入模型召回策略"""

import asyncio
import logging
import os
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import json
import httpx

try:
    import chromadb
    from chromadb.config import Settings
    CHROMADB_AVAILABLE = True
except ImportError:
    CHROMADB_AVAILABLE = False

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """文档数据类"""
    id: str
    content: str
    metadata: Optional[Dict[str, Any]] = None
    embedding: Optional[List[float]] = None


class SimpleRAG:
    """简单的RAG支持类"""
    
    def __init__(self, config: Dict[str, Any]):
        """初始化RAG系统"""
        self.config = config
        self.vector_store = None
        
        # Ollama嵌入模型配置
        ollama_config = self.config.get("ollama", {})
        self.ollama_base_url = ollama_config.get("base_url", "http://localhost:11434")
        self.ollama_model = ollama_config.get("model", "nomic-embed-text")
        
        # 初始化向量数据库
        if CHROMADB_AVAILABLE:
            self._init_vector_store()
        else:
            logger.warning("ChromaDB不可用，RAG功能将被禁用")
    
    def _init_vector_store(self):
        """初始化向量数据库"""
        try:
            vector_store_path = self.config.get(
                "vector_store_path", "data/db/rag_vector_store"
            )
            os.makedirs(vector_store_path, exist_ok=True)
            
            self.vector_store_client = chromadb.PersistentClient(
                path=vector_store_path,
                settings=Settings(allow_reset=True, anonymized_telemetry=False),
            )
            
            collection_name = self.config.get("collection_name", "stock_rag")
            self.vector_store = self.vector_store_client.get_or_create_collection(
                name=collection_name, metadata={"hnsw:space": "cosine"}
            )
            
            logger.info(f"向量数据库已初始化: {vector_store_path}")
        except Exception as e:
            logger.error(f"初始化向量数据库失败: {str(e)}")
            self.vector_store = None
    
    async def _get_ollama_embedding(self, text: str) -> List[float]:
        """获取Ollama嵌入"""
        try:
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    f"{self.ollama_base_url}/api/embeddings",
                    json={
                        "model": self.ollama_model,
                        "prompt": text
                    },
                    timeout=30.0
                )
                response.raise_for_status()
                result = response.json()
                return result["embedding"]
        except Exception as e:
            logger.error(f"获取Ollama嵌入失败: {str(e)}")
            raise
    
    async def add_documents(self, documents: List[Document]) -> bool:
        """添加文档到向量数据库，支持主流文本、pdf、doc、docx文件"""
        if not self.vector_store:
            logger.warning("向量数据库不可用，无法添加文档")
            return False


        import yaml
        import configparser
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            PdfReader = None
        try:
            import docx
        except ImportError:
            docx = None

        def extract_content(doc: Document) -> str:
            ext = os.path.splitext(doc.id)[1].lower()
            content = doc.content
            try:
                if ext == '.pdf' and PdfReader is not None:
                    with open(doc.id, 'rb') as f:
                        reader = PdfReader(f)
                        text = "\n".join(page.extract_text() or '' for page in reader.pages)
                    return text
                elif ext in {'.doc', '.docx'} and docx is not None:
                    d = docx.Document(doc.id)
                    text = "\n".join([p.text for p in d.paragraphs])
                    return text
                elif ext in {'.yaml', '.yml'}:
                    data = yaml.safe_load(content)
                    return yaml.dump(data, allow_unicode=True)
                elif ext in {'.ini', '.cfg'}:
                    parser = configparser.ConfigParser()
                    parser.read_string(content)
                    return str(dict(parser))
                elif ext in {'.csv', '.tsv'}:
                    return content  # 直接返回全部内容
                else:
                    return content  # 其它文本文件直接返回全部内容
            except Exception as e:
                logger.warning(f"解析文件 {doc.id} 失败: {e}")
                return content

        try:
            for doc in documents:
                doc.content = extract_content(doc)
                doc.embedding = await self._get_ollama_embedding(doc.content)
            self.vector_store.add(
                documents=[doc.content for doc in documents],
                embeddings=[doc.embedding for doc in documents],
                ids=[doc.id for doc in documents],
                metadatas=[doc.metadata for doc in documents]
            )
            logger.info(f"成功添加 {len(documents)} 个文档到向量数据库")
            return True
        except Exception as e:
            logger.error(f"添加文档失败: {str(e)}")
            return False
    
    async def retrieve(self, query: str, top_k: int = 5) -> List[Document]:
        """根据查询检索相关文档"""
        if not self.vector_store:
            logger.warning("向量数据库不可用，无法检索文档")
            return []
            
        try:
            # 生成查询嵌入
            query_embedding = await self._get_ollama_embedding(query)
            
            # 执行相似性搜索
            results = self.vector_store.query(
                query_embeddings=[query_embedding],
                n_results=top_k
            )
            
            # 构造返回结果
            documents = []
            for i in range(len(results['ids'][0])):
                doc = Document(
                    id=results['ids'][0][i],
                    content=results['documents'][0][i],
                    metadata=results['metadatas'][0][i] if results['metadatas'][0] else None,
                    embedding=results['embeddings'][0][i] if results['embeddings'] else None
                )
                documents.append(doc)
                
            logger.info(f"检索到 {len(documents)} 个相关文档")
            return documents
        except Exception as e:
            logger.error(f"文档检索失败: {str(e)}")
            return []

    async def get_all_documents(self) -> List[Document]:
        """获取所有文档"""
        if not self.vector_store:
            logger.warning("向量数据库不可用，无法获取文档")
            return []
            
        try:
            # 获取所有文档
            results = self.vector_store.get()
            
            # 构造返回结果
            documents = []
            for i in range(len(results['ids'])):
                doc = Document(
                    id=results['ids'][i],
                    content=results['documents'][i],
                    metadata=results['metadatas'][i] if results['metadatas'] else None,
                )
                documents.append(doc)
                
            logger.info(f"获取到 {len(documents)} 个文档")
            return documents
        except Exception as e:
            logger.error(f"获取文档失败: {str(e)}")
            return []


# 全局RAG实例
_rag_instance: Optional[SimpleRAG] = None


async def get_rag_instance(config: Dict[str, Any] = None) -> Optional[SimpleRAG]:
    """获取全局RAG实例（懒加载 + 单例）"""
    global _rag_instance
    if _rag_instance is not None:
        return _rag_instance
        
    if config is None:
        # 尝试从配置文件加载配置
        try:
            import yaml
            from pathlib import Path
            
            project_root = Path(__file__).resolve().parent.parent.parent
            settings_path = project_root / "config" / "settings.yaml"
            
            if settings_path.exists():
                with open(settings_path, "r", encoding="utf-8") as f:
                    settings = yaml.safe_load(f)
                    config = settings.get("rag", {})
            else:
                config = {}
        except Exception as e:
            logger.warning(f"加载RAG配置失败: {str(e)}")
            config = {}
    
    try:
        _rag_instance = SimpleRAG(config)
        return _rag_instance
    except Exception as e:
        logger.error(f"创建RAG实例失败: {str(e)}")
        return None